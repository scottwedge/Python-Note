# -*- coding: utf-8 -*-
"""
Created on 7/22/2019

@author: Jerry, Yiwen
"""
import os
import re
import math
import csv
from docx import Document
from docxcompose.composer import Composer
from NVUSdatabase import TS_DICT, SALES_DICT, PRICE_DICT
from process import process_dict

def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            found = False
            for run in p.runs:
                run.text, n = regex.subn(replace, run.text)
                found = found or (n != 0)
            if found is False:
                raise RuntimeWarning(
                    f'{regex.pattern} found in paragraph but not replaced.')
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)
def GetDesktopPath():
    return os.path.join(os.path.expanduser("~"), 'Desktop')

class quote(object):
    def __init__(self, info, price1, price2, BInumber, species, library_type, rRNAremoval_check):
        self.searchObj = [a.strip() for a in info.split("\t")]
        #process_dict
        self.process_dict = process_dict
        #client information
        self.firstname = self.searchObj[2]
        self.lastname = self.searchObj[3]
        self.var_client_email = str(self.searchObj[5])
        self.var_school = self.searchObj[4]
        self.var_client_name = self.searchObj[2] + ' ' + self.searchObj[3]
        #pre-input information
        self.price1 = price1
        self.price2 = price2
        self.BInumber = BInumber
        self.species = species
        self.rRNAremoval_check = rRNAremoval_check
        self.library_type = library_type
        #quote information
        self.service_name = ''
        self.var_library_type = ''
        self.quote_number = self.searchObj[1]
        self.var_quoteinfo = self.searchObj[6]
        self.seq_strategy = self.searchObj[15]
        try:
            self.data_output = int(self.searchObj[16])
        except:
            self.data_output = 0
        self.var_TAT = self.searchObj[18]
        self.var_quote = self.var_quoteinfo.split("-")
        self.quote_temp = self.searchObj[10]
        self.sam_num_1 = self.get_sam_num()[0]
        self.sam_num_2 = self.get_sam_num()[1]
        self.var_groupemail = 'Please input group email'
        #quote compose section
        self.filename_price = ""
        self.filename_requirement = ""
        self.filename_technicaltermsnotes = ""
        self.filename_bicontents = ""
        self.filename_submission = ""
        self.filename_quotationterms = ""
        self.composer_name = self.var_quoteinfo + ".docx"
    def product_check(self):
        '''get type of service'''
        if self.quote_temp.lower() in (s.lower() for s in PRICE_DICT.keys()):
            return self.quote_temp.lower()
        else:
            raise RuntimeError(
                f"product type {self.quote_temp} not found. Accepted products: " + ', ' .join(s for s in PRICE_DICT.keys()))
    def get_sam_num(self):
        '''get the sample number and analysis number'''
        if self.var_quote[0] == 'Davis' or self.var_quote[0] == 'AH' or self.var_quote[0] == 'ADH':
            self.sam_num_1 = int(self.var_quote[3])
            if self.var_quote[4] == 'premade':
                self.sam_num_2 = int(self.var_quote[5])
            else:
                self.sam_num_2 = self.sam_num_1
        else:
            self.sam_num_1 = int(self.var_quote[2])
            if self.var_quote[3] == 'premade':
                self.sam_num_2 = int(self.var_quote[4])
            else:
                self.sam_num_2 = self.sam_num_1
        return self.sam_num_1, self.sam_num_2
    def bio_type(self):
        self.bio_number = 0  # check the analysis type
        if 'WOBI' in self.var_quoteinfo:
            self.var_bio = 'Data Quality Control'
        elif 'quant' in self.var_quoteinfo.lower():
            self.var_bio = 'Standard quantification analysis'
        elif 'advance' in self.var_quoteinfo.lower():
            self.var_bio = 'Advance analysis'
        elif 'denovo' in self.var_quoteinfo.lower():
            self.var_bio = 'Denovo analysis'
        else:
            self.var_bio = 'Standard bioinformatics analysis'
        return self.var_bio
    def RIN_num(self):
        '''need add FFPE and other type'''
        if any(s.lower() in self.var_quoteinfo.split("-") for s in ['plant', 'flower', 'tree', 'fungi']):
            return 6.3
        elif any(s.lower() in self.var_quoteinfo.split("-") for s in ['ffpe', 'bacteria']):
            return 6.0
        else:
            return 6.8
    def get_species(self):  # check the species
        if not (list(filter(lambda x: x in self.quote_temp.lower(), ['premade','Metagenomics','amplicon']))):
            if 'human' in self.var_quoteinfo:
                self.species = 'Homo sapiens'
            elif 'mouse' in self.var_quoteinfo:
                self.species = 'Mus musculus'
            else:
                raise ValueError(
                    'species cannot be recognized, please input the latin name.')
        else:
            self.species = ''
        return self.species
    def people_check(self):  # '''check if the person in the database'''
        self.varTS = self.searchObj[13]
        self.varsales = self.searchObj[8]
        try:
            self.TS_name, self.TS_email, self.TS_OMS = TS_DICT[self.varTS.lower(
            ).strip()]
        except KeyError as e:
            raise KeyError("TS name " + str(e) + "is not found, valid TS names are " + ", ".join(
                s.capitalize() for s in TS_DICT.keys()))
        try:
            self.sales_name, self.sales_email = SALES_DICT[self.varsales]
        except KeyError as e:
            raise KeyError("sales name " + str(e) + " is not found, valid sales names are " + ", ".join(
                s.capitalize() for s in SALES_DICT.keys()))
        return self.TS_name, self.TS_email, self.TS_OMS, self.sales_name, self.sales_email
    def service_type(self):
        self.var_library_type = ""
        self.var_platform = ""
        self.var_UP_1 = 0
        self.var_UP_2 = 0
        if self.quote_temp.lower() == "eukmrnaseq":
            self.eukmrnaseq()
            self.rna_analysis()
        elif self.quote_temp.lower() == "prokmrnaseq":
            self.prokmrnaseq()
        elif self.quote_temp.lower() == "hwes":
            self.hwes()
        elif self.quote_temp.lower() == "mwes":
            self.mwes()
        elif self.quote_temp.lower() == "hwgs":
            self.hwgs()
        elif self.quote_temp.lower() == "pawgs":
            self.pawgs()
        elif self. quote_temp.lower() == "lncrnaseq":
            self.lncrna()
        elif self.quote_temp.lower() == "microbialwgs":
            self.bacterial_reseq()
        elif self.quote_temp.lower() == "metagenomics":
            self.metagenomics()
        elif "premade" in self.quote_temp.lower():
            self.premade_illumina()
        elif self.quote_temp.lower() == "amplicon":
            self.amplicon()
        elif self.quote_temp.lower() == 'circrnaseq':
            self.circRNA()
        elif self.quote_temp.lower() == 'srnaseq':
            self.smallRNA()
        elif list(filter(lambda x: x in self.quote_temp.lower(), ["isoseq", "pacbio"])):
            self.pacbio()
        else:
            raise RuntimeError(
                f'project type {self.quote_temp.lower()} not found.')
        print("filename: " + str(self.composer_name))
        print("var_UP_1: " + str(self.var_UP_1))
        print("var_UP_2: " + str(self.var_UP_2))
        print("data outtput: " + str(self.data_output))
        print("library type: " + str(self.var_library_type))
        print("platform: " + str(self.var_platform))
        return self.composer_name, self.var_UP_1, self.var_UP_2, self.data_output, self.var_library_type, self.var_platform
    def WOBI_check(self):
        if 'WOBI' in self.var_quoteinfo:
            self.var_UP_2 = 0
            self.filename_price = os.path.join(
                "./support-section", "Pricing_WOBI.docx")
            if self.var_quote[0] in ['AH', 'ADH']:
                self.filename_bicontents = os.path.join(
                    "./support-section", "WOBI_ADH.docx")
            else:
                self.filename_bicontents = os.path.join(
                    "./support-section", "WOBI_Novogene.docx")
        else:
            self.filename_price = os.path.join(
                "./support-section", "Pricing_WBI.docx")
            pass
    def ADH_check(self):
        if self.var_quote[0] in ['AH', 'ADH']:
            self.filename_submission = os.path.join(
                "./support-section", "Submission_ADH.docx")
        else:

            self.filename_submission = os.path.join(
                "./support-section", "Submission_Novogene.docx")  
    def rna_analysis(self):
        path = "./support-section/rna_product_analysis.csv"
        state = self.searchObj[11]
        globinclear = ''
        globinzero = ''
        ribo = ''
        stranded = ''
        if "globinzero" in self.var_quoteinfo.lower():
            globinzero = 1
            stranded = 1
        elif "globinclear" in self.var_quoteinfo.lower():
            globinclear = 1
        elif self.rRNAremoval_check == True:
            ribo = 1
            stranded = 1
        elif self.library_type == True:
            stranded = 1
        else:
            pass 
        value = [self.quote_number, ribo, globinclear, globinzero, stranded, self.sam_num_1, state]
        try:
            with open(path, 'a', newline = '') as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow(value)
            csvfile.close()   
            print("rna analysis worked")
        except:
            print("rna analysis didn't wroked")
    def eukmrnaseq(self):
        self.var_groupemail = 'us.rna@novogene.com'
        self.service_name = "Eukaryotic RNA-seq\n (lib prep & seq)"
        self.filename_quotationterms = os.path.join(
            "./support-section", "QuotationTerms.docx")
        self.filename_technicaltermsnotes = os.path.join(
            "./support-section", "TechnicalTermsNotes.docx")
        pwd = "./support-section/RNAproducts"
        pwd_BI = "./support-section/RNAproducts/Part4 BIcontents"
        #check ADH project
        self.ADH_check()
        #check the globin and stranded, prepare the lib&seq price
        if "globinzero" in self.var_quoteinfo.lower():
            self.process_dict["eukmrnaseq"][2][1] = 'LAA0020'
            self.filename_requirement = os.path.join(
                pwd, "Part2 SampleRequirements", "GlobinZero.docx")
            self.var_UP_1 = math.ceil(
                PRICE_DICT["EukmRNAseq"][7] + self.data_output * PRICE_DICT["EukmRNAseq"][1])
            self.var_library_type = 'stranded specific'
        elif "globinclear" in self.var_quoteinfo.lower() and (self.library_type == True):
            self.process_dict["eukmrnaseq"][2][1] = 'LAA0018'
            self.filename_requirement = os.path.join(
                pwd, "Part2 SampleRequirements", "GlobinClear.docx")
            self.var_UP_1 = math.ceil(
                PRICE_DICT["EukmRNAseq"][6] + self.data_output * PRICE_DICT["EukmRNAseq"][1])
            self.var_library_type = 'stranded specific'
        elif "globinclear" in self.var_quoteinfo.lower() and (self.library_type == False):
            self.process_dict["eukmrnaseq"][2][1] ='LAA0019'
            self.filename_requirement = os.path.join(
                pwd, "Part2 SampleRequirements", "GlobinClear.docx")
            self.var_UP_1 = math.ceil(
                PRICE_DICT["EukmRNAseq"][5] + self.data_output * PRICE_DICT["EukmRNAseq"][1])
            self.var_library_type = 'cDNA'
        elif self.rRNAremoval_check == True:
            self.process_dict["eukmrnaseq"][2][1] = 'LAA0044'
            self.filename_requirement = os.path.join(
                pwd, "Part2 SampleRequirements", "rRNAdepletion.docx")
            self.var_library_type = 'stranded specific'
            self.var_UP_1 = math.ceil(
                PRICE_DICT["ProkmRNAseq"][0] + self.data_output * PRICE_DICT["EukmRNAseq"][1])
        else:
            #check the stranded or not
            self.filename_requirement = os.path.join(
                pwd, "Part2 SampleRequirements", "EukmRNAseq.docx")
            if self.library_type == True:
                self.process_dict["eukmrnaseq"][2][1] = 'LAA0036'
                self.var_library_type = 'stranded specific'
                self.var_UP_1 = math.ceil(
                    PRICE_DICT["EukmRNAseq"][4] + self.data_output * PRICE_DICT["EukmRNAseq"][1])
            else:
                self.process_dict["eukmrnaseq"][2][1] = 'LAA0038'
                self.var_library_type = 'cDNA'
                #check the promotion price
                if self.data_output >= 6:
                    self.var_UP_1 = math.ceil(
                        PRICE_DICT["EukmRNAseq-promotion"][1]) if self.sam_num_1 < 24 else math.ceil(PRICE_DICT["EukmRNAseq-promotion"][0])
                    self.var_UP_1 = math.ceil(
                        self.var_UP_1 + (self.data_output - 6) * PRICE_DICT["EukmRNAseq"][1])
                else:
                    self.var_UP_1 = math.ceil(
                        PRICE_DICT["EukmRNAseq"][0] + self.data_output * PRICE_DICT["EukmRNAseq"][1])
        #check the BI and prepare the BI price
        self.WOBI_check()
        if "denovo" in self.var_quoteinfo.lower():
            if 'quant' in self.var_quoteinfo:
                self.process_dict["eukmrnaseq"][2][3] = 'LAA0099'
                self.var_UP_2 = math.ceil(PRICE_DICT["EukmRNAseq"][8])
                self.filename_bicontents = os.path.join(
                    pwd_BI, "quant_denovo.docx")
            else:
                self.process_dict["eukmrnaseq"][2][3] = 'LAA0120'
                self.var_UP_2 = math.ceil(PRICE_DICT["EukmRNAseq"][9])
                self.filename_bicontents = os.path.join(
                    pwd_BI, "standard_denovo.docx")
        elif 'WBI' in self.var_quoteinfo:
            if 'quant' in self.var_quoteinfo.lower():
                self.var_UP_2 = 0 if self.sam_num_1 >= 24 else math.ceil(
                    PRICE_DICT["EukmRNAseq"][2])
                self.var_UP_1 = math.ceil(
                    PRICE_DICT["EukmRNAseq-promotion"][1])
                if list(filter(lambda x: x in self.var_quoteinfo.lower(), ['human', 'mouse'])):
                    self.process_dict["eukmrnaseq"][2][3] = 'LAA0149'
                    self.filename_bicontents = os.path.join(
                        pwd_BI, "quant_medical.docx")
                else:
                    self.process_dict["eukmrnaseq"][2][3] = 'LAA0100'
                    self.filename_bicontents = os.path.join(
                        pwd_BI, "quant_agricultural.docx")
            elif list(filter(lambda x: x in self.var_quoteinfo.lower(), ['mapping', 'map'])):
                    self.filename_bicontents = os.path.join(
                        pwd_BI, "mapping_EukmRNAseq.docx")
            elif 'pdx' in self.var_quoteinfo.lower():
                    self.filename_bicontents = os.path.join(
                        pwd_BI, "PDX_EukmRNAseq.docx")
            else:
                self.var_UP_2 = math.ceil(
                    PRICE_DICT["EukmRNAseq"][3])
                if list(filter(lambda x: x in self.var_quoteinfo.lower(), ['human', 'mouse'])):
                    self.process_dict["eukmrnaseq"][2][3] = 'LAA0121'
                    self.filename_bicontents = os.path.join(
                        pwd_BI, "standard_medical.docx")
                else:
                    self.process_dict["eukmrnaseq"][2][3] = 'LAA0123'
                    self.filename_bicontents = os.path.join(
                        pwd_BI, "standard_agricultural.docx")
        else:
            self.process_dict["eukmrnaseq"][2][3] = 'LAA0136'
            pass
    def prokmrnaseq(self):
        self.var_groupemail = 'us.rna@novogene.com'
        pwd = "./support-section/RNAproducts"
        pwd_BI = "./support-section/RNAproducts/Part4 BIcontents"
        self.service_name = "Prokaryotic RNA-seq\n (lib prep & seq)"
        self.filename_quotationterms = os.path.join(
            "./support-section", "QuotationTerms.docx")
        self.filename_technicaltermsnotes = os.path.join(
            "./support-section", "TechnicalTermsNotes.docx")
        self.filename_requirement = os.path.join(
            pwd, "Part2 SampleRequirements", "ProkRNAseq.docx")
        self.var_UP_1 = math.ceil(
            PRICE_DICT["ProkmRNAseq"][0] + self.data_output * PRICE_DICT["ProkmRNAseq"][1])
        # ADH check
        self.ADH_check()
        #WBI process
        if list(filter(lambda x: x in self.var_quoteinfo.lower(), ['advance', 'denovo'])):
            self.process_dict["prokmrnaseq"][2][3] = "LAA0150"
            self.filename_bicontents = os.path.join(
                pwd_BI, "original_ProkRNAseq.docx")
            self.var_UP_2 = 0
        elif 'WBI' in self.var_quoteinfo:
            self.process_dict["prokmrnaseq"][2][3] = "LAA0152"
            self.filename_bicontents = os.path.join(
                pwd_BI, "standard_Prok.docx")
            self.var_UP_2 = PRICE_DICT["ProkmRNAseq"][2]
        else:
            self.process_dict["prokmrnaseq"][2][3] = "LAA0151"
            pass
        #check if this is WOBI project
        self.WOBI_check()
    def lncrna(self):
        self.var_groupemail = 'us.rna@novogene.com'
        pwd = "./support-section/RNAproducts"
        pwd_BI = "./support-section/RNAproducts/Part4 BIcontents"
        self.service_name = "Lnc RNA-seq\n (lib prep & seq)"
        self.filename_quotationterms = os.path.join(
            "./support-section", "QuotationTerms.docx")
        self.filename_technicaltermsnotes = os.path.join(
            "./support-section", "TechnicalTermsNotes.docx")
        self.filename_requirement = os.path.join(
            pwd, "Part2 SampleRequirements", "rRNAdepletion.docx")
        #ADH check
        self.ADH_check()
        self.var_UP_1 = math.ceil(
            PRICE_DICT["lncRNAseq"][0] + self.data_output * PRICE_DICT["lncRNAseq"][1])
        self.var_UP_2 = 0
        if "WBI" in self.var_quoteinfo:
            self.process_dict["lncrnaseq"][2][3] = "LAA0152"
            self.filename_bicontents = os.path.join(
                pwd_BI, "standard_lncRNA.docx")
            self.var_UP_2 = PRICE_DICT["lncRNAseq"][2]
        else:
            self.process_dict["lncrnaseq"][2][3] = "LAA0151"
            pass
        #WOBI check
        self.WOBI_check()
    def hwgs(self):
        self.var_groupemail = 'us.dna@novogene.com'
        pwd = "./support-section/DNAproducts"
        pwd_BI = "./support-section/DNAproducts/Part4 BIcontents"
        self.service_name = "Human WGS\n(lib prep & seq)"
        self.filename_quotationterms = os.path.join(
            "./support-section", "QuotationTerms.docx")
        self.filename_technicaltermsnotes = os.path.join(
            pwd, "Part3 TechnicalTerms", "DNA_TechnicalTerms.docx")
        if list(filter(lambda x: x in self.var_quoteinfo.lower(), ['pcr', 'pcr-free', 'pcrfree'])):
            self.process_dict["hwgs"][2][1] = "LAA0094"
            self.filename_requirement = os.path.join(
                pwd, "Part2 SampleRequirements", "PCRfree_requirement.docx")
        else:
            self.process_dict["hwgs"][2][1] = "LAA0029"
            self.filename_requirement = os.path.join(
                pwd, "Part2 SampleRequirements", "hwgs_requirement.docx")
        # lib prep and seq price check
        if self.data_output == 90:
            self.var_UP_1 = PRICE_DICT["hWGS"][0]
        else:
            self.var_UP_1 = math.ceil(PRICE_DICT["hWGS"][1] + PRICE_DICT["hWGS"][2] * self.data_output) if self.data_output > 90 else math.ceil(
                PRICE_DICT["hWGS"][1] + PRICE_DICT["hWGS"][3] * self.data_output)
        #check ADH
        self.ADH_check()
        if 'advance' in self.var_quoteinfo:
            self.process_dict["hwgs"][2][3] = "LAA0009"
            self.var_UP_2 = 0
            self.filename_bicontents = os.path.join(
                pwd_BI, "hwgs_advance.docx")
        elif 'WBI' in self.var_quoteinfo:
            self.process_dict["hwgs"][2][3] = "LAA0152"
            self.filename_bicontents = os.path.join(
                pwd_BI, "hwgs_standard.docx")
            self.var_UP_2 = PRICE_DICT["hWGS"][4]
        else:
            self.process_dict["hwgs"][2][3] = "LAA0151"
            self.filename_bicontents = os.path.join(
                pwd_BI, "hwgs_original.docx")    
        #check the BI
        self.WOBI_check()
    def hwes(self):
        self.var_groupemail = 'us.dna@novogene.com'
        pwd = "./support-section/DNAproducts"
        pwd_BI = "./support-section/DNAproducts/Part4 BIcontents"
        self.service_name = "Human WES\n(Exome Capture & lib prep & seq)"
        self.filename_quotationterms = os.path.join(
            "./support-section", "QuotationTerms.docx")
        self.filename_technicaltermsnotes = os.path.join(
            "./support-section", "TechnicalTermsNotes.docx")
        self.filename_requirement = os.path.join(
            pwd, "Part2 SampleRequirements", "wes_requirement.docx"
        )
        #lib prep and seq price check (include the FFPE check=> rRNAremoval check)
        if self.data_output == 6:
            self.var_UP_2 = PRICE_DICT["hWES"][4] if self.sam_num_1 < 24 else 0
            if self.rRNAremoval_check == True:
                self.var_UP_1 = PRICE_DICT["hWES"][7]
            else:
                self.var_UP_1 = PRICE_DICT["hWES"][0]
        elif self.data_output == 12:
            self.var_UP_2 = PRICE_DICT["hWES"][4] if self.sam_num_1 < 24 else 0
            if self.rRNAremoval_check == True:
                self.var_UP_1 = PRICE_DICT["hWES"][8]
            else:
                self.var_UP_1 = PRICE_DICT["hWES"][1]
        elif self.data_output > 12:
            self.var_UP_1 = math.ceil(
                PRICE_DICT["hWES"][6] + (self.data_output-6) * PRICE_DICT["hWES"][3])
            self.var_UP_2 = PRICE_DICT["hWES"][4] if self.sam_num_1 < 24 else 0
        elif self.data_output <= 6:
            self.var_UP_2 = math.ceil(PRICE_DICT["hWES"][4]) if (
                'WBI' in self.var_quoteinfo) else 0
            if self.rRNAremoval_check == True:
                self.var_UP_1 = PRICE_DICT["hWES"][9]
            else:
                self.var_UP_1 = math.ceil(
                    PRICE_DICT["hWES"][2] + self.data_output * PRICE_DICT["hWES"][3])
        else:
            raise RuntimeError(f'Please check quote name format.')
        # check the ADH
        self.ADH_check()
        # check the WOBI
        if list(filter(lambda x: x in self.var_quoteinfo.lower(), ['mapping', 'map', 'mappingonly', 'maponly'])):
            self.process_dict["hwes"][2][3] = "LAA0150"
            self.var_UP_2 = PRICE_DICT["hWES"][5]
            self.filename_bicontents = os.path.join(
                pwd_BI, "wes_standard.docx")
        elif 'advance' in self.var_quoteinfo.lower():
            self.process_dict["hwes"][2][3] = "LAA0157"
            self.filename_bicontents = os.path.join(
                pwd_BI, "wes_advance.docx"
            )
        elif 'WBI' in self.var_quoteinfo:
            self.process_dict["hwes"][2][3] = "LAA0152"
            self.filename_bicontents = os.path.join(
                pwd_BI, "wes_standard.docx")
            self.var_UP_2 = self.var_UP_2
        else:
            self.process_dict["hwes"][2][3] = "LAA0151"
            self.filename_bicontents = os.path.join(
                pwd_BI, "wes_original.docx")
        # check the WBI
        self.WOBI_check()
        if "ffpe" in self.var_quoteinfo.lower():
            self.process_dict["hwes"][2][1] = "LAA0063"
        else:
            self.process_dict["hwes"][2][1] = "LAA0075"
    def mwes(self):
        self.var_groupemail = 'us.dna@novogene.com'
        pwd = "./support-section/DNAproducts"
        pwd_BI = "./support-section/DNAproducts/Part4 BIcontents"
        self.service_name = "Mouse WES\n((Exome Capture & lib prep & seq)"
        self.filename_quotationterms = os.path.join(
            "./support-section", "QuotationTerms.docx")
        self.filename_technicaltermsnotes = os.path.join(
            pwd, "Part3 TechnicalTerms", "DNA_TechnicalTerms.docx")
        self.filename_requirement = os.path.join(
            pwd, "Part2 SampleRequirements", "wes_requirement.docx"
        )
        if self.data_output == 6:
            self.var_UP_1 = PRICE_DICT["mWES"][0]
            self.var_UP_2 = 50 if ("WBI" in self.var_quoteinfo) else 0
        elif self.data_output == 12:
            self.var_UP_1 = PRICE_DICT["mWES"][1]
            self.var_UP_2 = 50 if ("WBI" in self.var_quoteinfo) else 0
        elif self.data_output >= 6:
            self.var_UP_1 = math.ceil(
                PRICE_DICT["mWES"][2] + self.data_output * PRICE_DICT["mWES"][3])
            self.var_UP_2 = PRICE_DICT["mWES"][4] if (
                "WBI" in self.var_quoteinfo) else 0
        #check the ADH
        self.ADH_check()
        if 'advance' in self.var_quoteinfo.lower():
            self.process_dict["mwes"][2][3] = "LAA0009"
            self.filename_bicontents = os.path.join(
                pwd_BI, "wes_advance.docx"
            )
        elif 'WBI' in self.var_quoteinfo:
            self.process_dict["mwes"][2][3] = "LAA0152"
            self.filename_bicontents = os.path.join(
                pwd_BI, "wes_standard.docx")
            self.var_UP_2 = self.var_UP_2
        else:
            self.process_dict["mwes"][2][3] = "LAA0151"
            self.filename_bicontents = os.path.join(
                pwd_BI, "wes_original.docx")
        #Check the WOBI
        self.WOBI_check()
    def pawgs(self):
        self.var_groupemail = 'us.dna@novogene.com'
        pwd = "./support-section/DNAproducts"
        pwd_BI = "./support-section/DNAproducts/Part4 BIcontents"
        self.service_name = "Animal and Plant WGS\n(lib prep & seq)"
        self.filename_quotationterms = os.path.join(
            "./support-section", "QuotationTerms.docx")
        self.filename_technicaltermsnotes = os.path.join(
            "./support-section", "TechnicalTermsNotes.docx")
        if list(filter(lambda x: x in self.var_quoteinfo.lower(), ['pcr', 'pcr-free', 'pcrfree'])):
            self.process_dict["pawgs"][2][1] = "LAA0094"
            self.filename_requirement = os.path.join(
                pwd, "Part2 SampleRequirements", "PCRfree_requirement.docx")
        else:
            self.process_dict["pawgs"][2][1] = "LAA0029"
            self.filename_requirement = os.path.join(
                pwd, "Part2 SampleRequirements", "pawgs_requirement.docx")
        self.var_UP_1 = math.ceil(
            PRICE_DICT["PAWGS"][0] + self.data_output * PRICE_DICT["PAWGS"][1])
        self.var_UP_2 = 0
        #check the ADH
        self.ADH_check()
        if list(filter(lambda x: x in self.var_quoteinfo.lower(), ['wbi-std', 'std', 'wbi'])):
            self.process_dict["pawgs"][1][3] = "GY0004"
            self.process_dict["pawgs"][2][3] = "LAA0131"
            self.var_UP_2 = PRICE_DICT["PAWGS"][2] * self.data_output
            self.filename_bicontents = os.path.join(
                pwd_BI, "pawgs_standard.docx")
        elif list(filter(lambda x: x in self.var_quoteinfo.lower(), ['advance', 'std-adv', 'adv'])):
            self.process_dict["pawgs"][1][3] = "GY0004"
            self.process_dict["pawgs"][2][3] = "LAA0010"
            self.filename_bicontents = os.path.join(
                pwd_BI, "pawgs_advance.docx")
            self.var_UP_2 = PRICE_DICT["PAWGS"][3] * self.data_output
        else:
            self.process_dict["pawgs"][1][3] = "GY0005"
            self.process_dict["pawgs"][2][3] = "LAA0151"
            self.filename_bicontents = os.path.join(
                pwd_BI, "pawgs_original.docx")
        #check the WOBI
        self.WOBI_check()
    def bacterial_reseq(self):
        self.var_groupemail = 'us.dna@novogene.com'
        pwd = "./support-section/microbial"
        pwd_BI = "./support-section/microbial/Part4 BIcontents"
        self.service_name = "Microbial Resequencing (WGS)\n(lib prep & seq)"
        self.filename_quotationterms = os.path.join(
            "./support-section", "QuotationTerms.docx")
        self.filename_technicaltermsnotes = os.path.join(
            "./support-section", "TechnicalTermsNotes.docx")
        self.filename_requirement = os.path.join(
            pwd, "Part2 SampleRequirements", "microbial_bacteria_requirement.docx")
        self.var_UP_1 = math.ceil(
            PRICE_DICT["MicrobialWGS"][0] + self.data_output * PRICE_DICT["MicrobialWGS"][1])
        #ADH check
        self.ADH_check()
        #WBI process
        if 'WBI' in self.var_quoteinfo:
            self.process_dict["microbialwgs"][2][3] = "LAA0160"
            self.var_UP_2 = PRICE_DICT["MicrobialWGS"][2]
            if 'bacteria' in self.var_quoteinfo.lower():
                self.filename_bicontents = os.path.join(
                    pwd_BI, "standard_microbial_bacteria.docx")
            elif 'fungi' in self.var_quoteinfo.lower():
                self.filename_bicontents = os.path.join(
                    pwd_BI, "standard_microbial_fungal.docx")
            else:
                self.filename_bicontents = os.path.join(
                    pwd_BI, "original_microbial.docx")
        else:
            self.process_dict["microbialwgs"][2][3] = "LAA0151"
            pass
        #WOBI check
        self.WOBI_check()
        if "adv" in self.var_quoteinfo.lower():
            self.process_dict["microbialwgs"][2][3] = "LAA0009"
    def metagenomics(self):
        self.var_groupemail = 'us.dna@novogene.com'
        pwd = "./support-section/microbial"
        pwd_BI = "./support-section/microbial/Part4 BIcontents"
        self.service_name = "Metagenomics\n(lib prep & seq)"
        self.filename_quotationterms = os.path.join(
            "./support-section", "QuotationTerms.docx")
        self.filename_technicaltermsnotes = os.path.join(
            "./support-section", "TechnicalTermsNotes.docx")
        self.filename_requirement = os.path.join(
            pwd, "Part2 SampleRequirements", "meta_requirement.docx"
        )
        self.var_UP_2 = 0
        if self.data_output == 6:
            self.var_UP_1 = math.ceil(PRICE_DICT["Metagenomics"][0]) if self.sam_num_1 < 48 else math.ceil(PRICE_DICT["Metagenomics"][5])
        elif self.data_output == 12:
            self.var_UP_1 = math.ceil(PRICE_DICT["Metagenomics"][1])
        else:
            self.var_UP_1 = math.ceil(PRICE_DICT["Metagenomics"][2] + self.data_output * PRICE_DICT["Metagenomics"][4])
        #check ADH
        self.ADH_check()
        #check WBI
        if "WBI" in self.var_quoteinfo:
            self.process_dict["metagenomics"][1][3] = "GY0004"
            self.process_dict["metagenomics"][2][3] = "LAA0143"
            self.var_UP_2 = math.ceil(
                PRICE_DICT["Metagenomics"][3] * self.data_output)
            self.filename_bicontents = os.path.join(
                pwd_BI, "standard_meta.docx")
        else:
            self.process_dict["metagenomics"][1][3] = "GY0005"
            self.process_dict["metagenomics"][2][3] = "LAA0151"
            pass
        #check WOBI
        self.WOBI_check()
    def premade_illumina(self):
        self.var_groupemail = 'us.library@novogene.com'
        pwd = "./support-section/premade"
        self.filename_quotationterms = os.path.join(
            "./support-section", "QuotationTerms.docx")
        self.filename_technicaltermsnotes = os.path.join(
            "./support-section", "premade", "Part3 TechnicalTerms", "PremadeTechnicalTerms.docx")
        self.filename_price = os.path.join(
            "./support-section", "premade", "Part1 Pricing", "Pricing_premade.docx")
        self.filename_requirement = os.path.join(
            pwd, "Part2 SampleRequirements", "Premade_Requirements.docx"
        )
        #ADH check
        self.ADH_check()
        if self.var_quote[0] in ['AH','ADH']:
            self.filename_bicontents = os.path.join(
                "./support-section", "WOBI_ADH.docx")
        else:
            self.filename_bicontents = os.path.join(
                "./support-section", "WOBI_Novogene.docx")
        self.sam_num_2 = self.get_sam_num()[1]
        self.var_UP_1 = PRICE_DICT["Premade-HiSeq"][0]
        if "flowcell" in self.var_quoteinfo.lower():
            self.data_output = 'flow cell'
        else:
            self.data_output = 'lane'
        if "novas4" in self.quote_temp.lower():
            self.var_platform = 'NovaseqS4'
            if self.sam_num_2 <= 3:
                self.var_UP_2 = PRICE_DICT["Premade-NovaS4"][1]
            elif self.sam_num_2 > 3 and self.sam_num_2 < 6:
                self.var_UP_2 = PRICE_DICT["Premade-NovaS4"][2]
            elif self.sam_num_2 > 5 and self.sam_num_2 < 8:
                self.var_UP_2 = PRICE_DICT["Premade-NovaS4"][3]
            else:
                self.var_UP_2 = PRICE_DICT["Premade-NovaS4"][4]
        elif "novasp" in self.quote_temp.lower():
            self.var_platform = 'NovaseqSP'
            if "PE250" in self.seq_strategy.lower():
                self.var_UP_2 = PRICE_DICT["Premade-NovaSP"][1]
            elif list(filter(lambda x: x in self.seq_strategy.lower(), ["PE50", "SE50"])):
                self.var_UP_2 = PRICE_DICT["Premade-NovaSP"][3]
            else:
                RuntimeError

        else:
            self.var_UP_2 = PRICE_DICT["Premade-HiSeq"][1]
            self.var_platform = 'Hiseq Platform'
    def amplicon(self):
        self.var_groupemail = 'us.dna@novogene.com'
        pwd = "./support-section/microbial"
        pwd_BI = "./support-section/microbial/Part4 BIcontents"
        self.service_name = "Amplicon (lib prep & seq)"
        self.filename_quotationterms = os.path.join(
            "./support-section", "QuotationTerms.docx")
        self.filename_technicaltermsnotes = os.path.join(
            "./support-section", "TechnicalTermsNotes.docx")
        self.filename_price = os.path.join(
            pwd, "Part1 Pricing", "Pricing_amplicon.docx")
        self.filename_requirement = os.path.join(
            pwd, "Part2 SampleRequirements", "amplicon_requirement.docx"
        )
        #ADH check
        self.ADH_check()
        #price and WBI later.......
        if "WBI" in self.var_quoteinfo:
            self.filename_bicontents = os.path.join(
                pwd_BI, "standard_amplicon.docx")
        elif "advance" in self.var_quoteinfo:
            self.filename_bicontents = os.path.join(
                pwd_BI, "advance_amplicon.docx")
        else:
            self.filename_bicontents = os.path.join(
                pwd_BI, "original_amplicon.docx")
        #check WOBI
        self.WOBI_check()
    def circRNA(self):
        self.var_groupemail = 'us.rna@novogene.com'
        self.service_name = "CircRNA"
        pwd = "./support-section/RNAproducts"
        self.filename_quotationterms = os.path.join(
            "./support-section", "QuotationTerms.docx")
        self.filename_price = os.path.join(
            pwd, "CircRNAtemplate.docx")
        #ADH check
        self.ADH_check()
    def smallRNA(self):
        self.var_groupemail = 'us.rna@novogene.com'
        self.service_name = "Small RNA"
        pwd = "./support-section/RNAproducts"
        self.filename_quotationterms = os.path.join(
            "./support-section", "QuotationTerms.docx")
        self.filename_price = os.path.join(
            pwd, "smallRNAtemplate.docx")
        #ADH check
        self.ADH_check()
    def pacbio(self):
        pwd = "./support-section/pacbio"
        pwd_BI = "./support-section/pacbio/Part4 BIcontents"
        self.var_groupemail = 'us.dna@novogene.com'
        self.filename_quotationterms = os.path.join(
            "./support-section", "QuotationTerms.docx")
        self.filename_technicaltermsnotes = os.path.join(
            pwd, "Part3 TechnicalTerms", "Pacbio_TechnicalTermsNotes.docx")
        if "isoseq" in self.quote_temp.lower():
            self.filename_price = os.path.join(
                pwd, "Part1 Pricing", "RNA_pricing.docx")
            self.filename_requirement = os.path.join(
                pwd, "Part2 SampleRequirements", "RNA_requirement.docx"
            )
            self.filename_bicontents = os.path.join(
                pwd_BI, "pacbio_RNA_Standard.docx")
        else:
            self.filename_price = os.path.join(
                pwd, "Part1 Pricing", "DNA_pricing.docx")
            self.filename_requirement = os.path.join(
                pwd, "Part2 SampleRequirements", "DNA_requirement.docx"
            )
            self.filename_bicontents = os.path.join(
                pwd_BI, "pacbio_DNA_Standard.docx")
        #ADH check
        self.ADH_check()
        #check WOBI
        if 'WOBI' in self.var_quoteinfo:
            self.var_UP_2 = 0
            if self.var_quote[0] in ['AH','ADH']:
                self.filename_bicontents = os.path.join(
                    "./support-section", "WOBI_ADH.docx")
            else:
                self.filename_bicontents = os.path.join(
                    "./support-section", "WOBI_Novogene.docx")
        else:
            pass
    def compose(self):
        pwd = GetDesktopPath()
        outpath = os.path.join(pwd, self.var_quoteinfo + '.docx')
        master = Document(self.filename_price)
        composer = Composer(master)
        for doc in [self.filename_requirement, self.filename_technicaltermsnotes, self.filename_bicontents, self.filename_submission, self.filename_quotationterms]:
            if doc:
                composer.append(Document(doc))
        self.composer_name = outpath
        composer.save(self.composer_name)
