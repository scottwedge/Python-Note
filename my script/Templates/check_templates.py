# -*- coding: utf-8 -*-
"""
Script to check if a template has "broken" key string

@author: Yiwen Zhang
"""

import os
import argparse
from docx import Document


def check_docx_var(filename):
    """ Function to check if a docx file has broken key strings
    like 'var', '_client_name'
    """
    # the below names set should be consistent with info_dict defined in
    # Quote.generate_quote function
    var_names = {
        'VAR_CLIENT_NAME',
        'VAR_CLIENT_EMAIL',
        'VAR_CLIENT_SCHOOL',
        'VAR_QUOTE_NUM',
        'VAR_ALY',
        'VAR_RIN',
        'VAR_TAT',
        'VAR_UP_1',
        'VAR_UP_2',
        'VAR_SPECIES',
        'VAR_TS',
        'TS_EMAIL',
        'VAR_SALES',
        'SALES_EMAIL',
        'TOTAL_1',
        'TOTAL_2',
        'TOTAL_FINAL',
        'SAM_NUM_1',
        'SAM_NUM_2',
        'DATA_OUTPUT',
        'VAR_LIBRARY_TYPE',
        'VAR_PLATFORM',
    }
    doc = Document(filename)
    for p in doc.paragraphs:
        for v in var_names:
            # count the total number that var name appears
            count_total = p.text.count(v)
            if count_total > 0:
                # count the number that var name is found in each "run"
                count_valid = sum(run.text.count(v) for run in p.runs)
                if count_valid != count_total:
                    print(f"Found broken key {v} in {filename}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for v in var_names:
                        count_total = p.text.count(v)
                        if count_total > 0:
                            # count the number that var name is found in each "run"
                            count_valid = sum(run.text.count(v)
                                              for run in p.runs)
                            if count_valid != count_total:
                                print(f"Found broken key {v} in {filename}")


def main():
    parser = argparse.ArgumentParser(description="Docx Template check")
    parser.add_argument(
        'infiles', nargs='*', help='docx files to check, default all .docx files in current directory')
    args = parser.parse_args()
    # add all .docx files in
    infiles = args.infiles
    if not infiles:
        infiles = [f for f in os.listdir(
            '.') if os.path.splitext(f)[-1] == '.docx']
    if not infiles:
        print("No .docx files has been found")
        return
    for f in infiles:
        check_docx_var(f)
    print("All checks finished")


if __name__ == '__main__':
    main()
